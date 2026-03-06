from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def _status_zh(status: Any) -> str:
    mapping = {
        "non_compliant": "不符合",
        "risk": "风险",
        "needs_manual": "需人工复核",
    }
    return mapping.get(str(status or "").strip(), str(status or "").strip())


def write_docx_report(report: dict[str, Any], output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    out = output_dir / "review_report.docx"

    doc = Document()
    doc.add_heading("标书审查报告", level=1)

    summary = report.get("summary", {})
    doc.add_heading("概览", level=2)
    p = doc.add_paragraph()
    p.add_run(f"硬性要求数: {summary.get('requirement_count', 0)}\n")
    p.add_run(f"不符合项数: {summary.get('non_compliant_count', 0)}\n")
    p.add_run(f"风险项数: {summary.get('risk_count', 0)}\n")
    p.add_run(f"需人工复核数: {summary.get('needs_manual_count', 0)}\n")
    p.add_run(f"总发现数: {summary.get('finding_count', 0)}")

    findings = report.get("findings", []) or []
    doc.add_heading("发现项明细", level=2)
    f_table = doc.add_table(rows=1, cols=7)
    f_table.style = "Table Grid"
    headers = ["ID", "条款ID", "结论", "问题", "招标证据", "投标证据", "建议"]
    for i, h in enumerate(headers):
        f_table.cell(0, i).text = h
    for finding in findings:
        row = f_table.add_row().cells
        row[0].text = str(finding.get("id", ""))
        row[1].text = str(finding.get("requirement_id", ""))
        row[2].text = _status_zh(finding.get("status", ""))
        row[3].text = str(finding.get("issue", ""))
        row[4].text = str(finding.get("tender_evidence", ""))
        row[5].text = str(finding.get("bid_evidence", ""))
        row[6].text = str(finding.get("recommendation", ""))

    requirements = report.get("requirements", []) or []
    doc.add_heading("硬性要求清单", level=2)
    r_table = doc.add_table(rows=1, cols=4)
    r_table.style = "Table Grid"
    r_headers = ["条款ID", "类别", "内容", "来源"]
    for i, h in enumerate(r_headers):
        r_table.cell(0, i).text = h
    for req in requirements:
        row = r_table.add_row().cells
        row[0].text = str(req.get("id", ""))
        row[1].text = str(req.get("category", ""))
        row[2].text = str(req.get("text", ""))
        row[3].text = str(req.get("source", ""))

    doc.save(out)
    return out

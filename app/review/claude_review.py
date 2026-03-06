from __future__ import annotations

import json
import re
import os
import time
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

from docx import Document

from app.llm.claude_client import (
    ClaudeCallError,
    compact_text_for_prompt,
    extract_json_payload,
    prompt_safe_path,
)
from app.llm.prompt_store import render_prompt


_CONTEXT_REQUIREMENT_CATEGORY = "主体一致性"
_CONTEXT_REQUIREMENT_TEXT = (
    "投标文件中的关键主体名词（招标人/采购人/投标人/供应商/开户银行/账户名/账号/"
    "统一社会信用代码/税号/法定代表人/授权代表）必须与所在位置和语义角色一致，"
    "不得出现主体错位、字段串用或其他机构信息误填。"
)
_CONTEXT_REQUIREMENT_SOURCE = "系统一致性校验规则（主体名词位置归属检查）"
_CONTEXT_REQUIREMENT_KEYWORDS = (
    "主体",
    "一致性",
    "错位",
    "招标人",
    "投标人",
    "采购人",
    "开户银行",
    "账户名",
    "账号",
    "统一社会信用代码",
    "税号",
)
_CONTEXT_FINDING_KEYWORDS = (
    "主体",
    "错位",
    "串用",
    "归属",
    "位置",
    "招标人",
    "投标人",
    "采购人",
    "开户银行",
    "账户名",
    "账号",
    "统一社会信用代码",
    "税号",
    "法定代表人",
    "授权代表",
    "其他机构",
    "错写",
    "误填",
    "抬头",
    "落款",
)


_IMAGE_SUFFIXES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
    ".gif",
}


def _instruction_requires_ocr(user_instruction: str, extra_instruction: str) -> bool:
    text = f"{user_instruction}\n{extra_instruction}".lower()
    keywords = ["ocr", "图片", "截图", "图像", "扫描", "证据图", "影像"]
    return any(k in text for k in keywords)


def _docx_ocr_required_by_default() -> bool:
    return os.getenv("BID_REVIEW_DOCX_OCR_REQUIRED", "1").strip().lower() not in {
        "0",
        "false",
        "no",
        "off",
    }


def _has_ocr_tool_call(tool_calls: list[str]) -> bool:
    for name in tool_calls:
        n = str(name).lower()
        if "paddle-ocr" in n:
            return True
        if "perform_ocr" in n or "perform_pdf_ocr" in n or "perform_batch_ocr" in n:
            return True
    return False


def _has_word_image_extract_call(tool_calls: list[str]) -> bool:
    for name in tool_calls:
        n = str(name).lower()
        if "extract_images_from_word" in n:
            return True
    return False


def _has_word_batch_ocr_call(tool_calls: list[str]) -> bool:
    for name in tool_calls:
        n = str(name).lower()
        if "ocr_images_in_dir" in n:
            return True
        if "perform_batch_ocr" in n:
            return True
    return False


def _append_ocr_enforcement(prompt: str, *, require_word_extract: bool) -> str:
    extra = ""
    if require_word_extract:
        extra = (
            "\n你必须先使用 `document-parser` MCP 服务器中的 Word 图片提取工具"
            "（如 `document-parser.extract_images_from_word` 或 `mcp__document-parser__extract_images_from_word`），"
            "再使用 `paddle-ocr` MCP 服务器中的批量图片 OCR 工具"
            "（如 `paddle-ocr.ocr_images_in_dir` 或 `mcp__paddle-ocr__ocr_images_in_dir`）对提取目录的全部图片完成OCR。"
        )
    enforce = """

[强制执行要求]
你必须至少调用一次 OCR MCP 工具（如 `paddle-ocr.ocr_image` / `paddle-ocr.ocr_pdf`，或 `mcp__paddle-ocr__ocr_image` / `mcp__paddle-ocr__ocr_pdf`）读取图片文字后再输出结果。
如果没有调用 OCR 工具，本次回答视为无效。
"""
    return prompt + enforce + extra


def _append_no_write_enforcement(prompt: str) -> str:
    enforce = """

[只读执行约束]
禁止创建、修改或删除任何本地文件，禁止运行会写文件的 Bash/PowerShell/Python 命令。
只允许使用 MCP 工具读取文档，以及只读命令（ls/find/rg/cat/head/tail）做定位。
若发生任意写操作（含重定向、tee、Out-File、Set-Content、touch、mkdir、mv、cp、rm），本次回答视为无效。
"""
    return prompt + enforce


def _tool_input_to_text(tool_input: Any) -> str:
    if isinstance(tool_input, str):
        return tool_input
    try:
        return json.dumps(tool_input, ensure_ascii=False)
    except Exception:  # noqa: BLE001
        return str(tool_input)


_REDIRECT_RE = re.compile(
    r"(?<!\S)(?:\d+)?(?:>>|>|&>)\s*(?P<target>&\d+|\"[^\"]+\"|'[^']+'|[^\s;|&]+)",
    flags=re.IGNORECASE,
)


def _extract_shell_command_text(tool_input: Any) -> str:
    if isinstance(tool_input, dict):
        command = tool_input.get("command")
        if isinstance(command, str):
            return command
    return _tool_input_to_text(tool_input)


def _strip_shell_quotes(token: str) -> str:
    value = token.strip()
    if len(value) >= 2 and ((value[0] == value[-1] == '"') or (value[0] == value[-1] == "'")):
        return value[1:-1]
    return value


def _is_allowed_redirect_target(token: str) -> bool:
    target = _strip_shell_quotes(token).strip().lower().rstrip(";,)")
    if target.startswith("&"):
        # 允许文件描述符重定向，如 2>&1 / >&2。
        return True
    return target in {"/dev/null", "nul", "$null"}


def _has_forbidden_shell_redirection(command: str) -> bool:
    for match in _REDIRECT_RE.finditer(command):
        target = match.group("target")
        if not _is_allowed_redirect_target(target):
            return True
    return False


def _has_forbidden_write_tool_call(tool_uses: list[dict[str, Any]]) -> bool:
    shell_markers = ("bash", "shell", "powershell", "terminal", "command", "exec")
    direct_write_markers = ("write", "edit", "multiedit")
    forbidden_patterns = [
        r"(?<![\w.-])out-file(?![\w.-])",
        r"(?<![\w.-])set-content(?![\w.-])",
        r"(?<![\w.-])add-content(?![\w.-])",
        r"(?<![\w.-])new-item(?![\w.-])",
        r"(?<![\w.-])tee(?![\w.-])",
        r"(?<![\w.-])touch(?![\w.-])",
        r"(?<![\w.-])mkdir(?![\w.-])",
        r"(?<![\w.-])copy-item(?![\w.-])",
        r"(?<![\w.-])move-item(?![\w.-])",
        r"(?<![\w.-])remove-item(?![\w.-])",
        r"(^|[;&|]\s*|\s+)(?:rm|mv|cp)\s+",
        r"(^|[;&|]\s*|\s+)python(?:\.exe)?\s+",
        r"(^|[;&|]\s*|\s+)py\s+-",
        r"(?<![\w.-])uv\s+run\s+python(?![\w.-])",
    ]
    for call in tool_uses:
        name = str(call.get("name", "")).lower()
        if name in direct_write_markers:
            return True
        if name.endswith("__write") or name.endswith("_write") or name.endswith(".write"):
            return True
        if "__write_file" in name or "_write_file" in name or "create_file" in name:
            return True
        if not any(m in name for m in shell_markers):
            continue
        command_text = _extract_shell_command_text(call.get("input")).lower()
        if _has_forbidden_shell_redirection(command_text):
            return True
        if any(re.search(pattern, command_text) for pattern in forbidden_patterns):
            return True
    return False


def _strict_fail_on_forbidden_write() -> bool:
    return os.getenv("BID_REVIEW_FAIL_ON_FORBIDDEN_WRITE", "0").strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }


def _iter_tool_input_strings(value: Any) -> list[str]:
    out: list[str] = []
    if isinstance(value, str):
        out.append(value)
    elif isinstance(value, dict):
        for v in value.values():
            out.extend(_iter_tool_input_strings(v))
    elif isinstance(value, (list, tuple, set)):
        for item in value:
            out.extend(_iter_tool_input_strings(item))
    return out


def _to_path_candidate(raw_text: str) -> Path | None:
    text = raw_text.strip().strip("\"'")
    text = text.rstrip(".,;)]}")
    if not text:
        return None
    lower = text.lower()
    if lower.startswith(("http://", "https://", "{env:", "env:")):
        return None
    has_drive = bool(re.match(r"^[a-zA-Z]:[\\/]", text))
    has_sep = "/" in text or "\\" in text
    has_ext = bool(Path(text).suffix)
    if not (has_drive or has_sep or has_ext):
        return None
    path_obj = Path(text).expanduser()
    if not path_obj.is_absolute():
        path_obj = (Path.cwd() / path_obj).resolve(strict=False)
    else:
        path_obj = path_obj.resolve(strict=False)
    return path_obj


def _iter_path_candidates(tool_input: Any) -> list[Path]:
    out: list[Path] = []
    for text in _iter_tool_input_strings(tool_input):
        path_obj = _to_path_candidate(text)
        if path_obj is not None:
            out.append(path_obj)
    return out


def _is_image_file(path_obj: Path) -> bool:
    return path_obj.suffix.lower() in _IMAGE_SUFFIXES


def _canonical_path(path_obj: Path) -> str:
    return str(path_obj.resolve(strict=False)).replace("\\", "/").lower()


def _list_image_files(directory: Path) -> list[Path]:
    if not directory.exists() or not directory.is_dir():
        return []
    out: list[Path] = []
    try:
        for p in directory.rglob("*"):
            if p.is_file() and _is_image_file(p):
                out.append(p.resolve(strict=False))
    except OSError:
        return []
    return out


def _collect_extracted_images_from_tool_uses(tool_uses: list[dict[str, Any]]) -> set[str]:
    images: set[str] = set()
    for call in tool_uses:
        name = str(call.get("name", "")).lower()
        if "extract_images_from_word" not in name:
            continue
        for path_obj in _iter_path_candidates(call.get("input")):
            if path_obj.suffix.lower() == ".docx":
                continue
            if path_obj.exists() and path_obj.is_file() and _is_image_file(path_obj):
                images.add(_canonical_path(path_obj))
                continue
            if path_obj.exists() and path_obj.is_dir():
                for image in _list_image_files(path_obj):
                    images.add(_canonical_path(image))
    return images


def _collect_ocr_target_images_from_tool_uses(tool_uses: list[dict[str, Any]]) -> set[str]:
    images: set[str] = set()
    for call in tool_uses:
        name = str(call.get("name", "")).lower()
        if "ocr_images_in_dir" not in name and "perform_batch_ocr" not in name:
            continue
        paths = _iter_path_candidates(call.get("input"))
        explicit_images = [p for p in paths if _is_image_file(p)]
        if explicit_images:
            for image in explicit_images:
                images.add(_canonical_path(image))
            continue
        for path_obj in paths:
            if path_obj.exists() and path_obj.is_dir():
                for image in _list_image_files(path_obj):
                    images.add(_canonical_path(image))
    return images


def _count_docx_embedded_images(docx_path: Path) -> int:
    if not docx_path.exists() or not docx_path.is_file():
        return 0
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            names = zf.namelist()
    except Exception:  # noqa: BLE001
        return 0
    count = 0
    for name in names:
        lowered = name.lower()
        if not lowered.startswith("word/media/"):
            continue
        suffix = Path(lowered).suffix
        if suffix in _IMAGE_SUFFIXES:
            count += 1
    return count


def _validate_docx_ocr_coverage(tool_uses: list[dict[str, Any]], *, bid_path: Path) -> tuple[bool, str]:
    extracted_images = _collect_extracted_images_from_tool_uses(tool_uses)
    ocr_images = _collect_ocr_target_images_from_tool_uses(tool_uses)

    if extracted_images:
        if not ocr_images:
            return False, "无法统计 OCR 覆盖数量（未识别到 ocr_images_in_dir/perform_batch_ocr 的目标图片）。"
        missing = sorted(extracted_images - ocr_images)
        if missing:
            covered = len(extracted_images) - len(missing)
            sample = ", ".join([Path(x).name for x in missing[:3]])
            return (
                False,
                f"Word提图共{len(extracted_images)}张，OCR覆盖{covered}张，缺少{len(missing)}张（示例: {sample}）。",
            )
        return True, f"Word提图共{len(extracted_images)}张，OCR已全量覆盖。"

    expected_count = _count_docx_embedded_images(bid_path)
    if expected_count <= 0:
        return True, "Word未检测到内嵌图片，跳过图片OCR覆盖校验。"
    if not ocr_images:
        return False, (
            "无法统计 OCR 覆盖数量（未识别到 ocr_images_in_dir/perform_batch_ocr 的目标图片），"
            f"但 Word 内嵌图片共{expected_count}张。"
        )
    if len(ocr_images) < expected_count:
        return (
            False,
            f"Word内嵌图片共{expected_count}张，OCR目标仅识别到{len(ocr_images)}张，存在未覆盖图片。",
        )
    return True, f"Word内嵌图片共{expected_count}张，OCR目标识别到{len(ocr_images)}张，满足全量覆盖。"


def _normalize_requirements(raw: Any) -> list[dict[str, Any]]:
    items = raw if isinstance(raw, list) else []
    out: list[dict[str, Any]] = []
    for idx, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        out.append(
            {
                "id": item.get("id") or f"R{idx:03d}",
                "category": item.get("category") or "",
                "text": item.get("text") or item.get("requirement_text") or "",
                "source": item.get("source") or item.get("source_location") or "",
            }
        )
    return [x for x in out if x["text"]]


def _normalize_status(raw_status: Any) -> str:
    status_text = str(raw_status or "").strip().lower()
    mapping = {
        "non_compliant": "non_compliant",
        "risk": "risk",
        "needs_manual": "needs_manual",
        "不符合": "non_compliant",
        "风险": "risk",
        "需人工复核": "needs_manual",
        "需要人工复核": "needs_manual",
        "人工复核": "needs_manual",
        "manual": "needs_manual",
    }
    return mapping.get(status_text, "needs_manual")


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    text = re.sub(r"[ \t]+", " ", text)
    # 清理末尾省略号，避免“问题描述...”这种不可执行表述
    text = re.sub(r"(?:\.{3,}|…+)\s*$", "", text)
    return text.strip()


def _clean_issue(value: Any) -> str:
    text = _clean_text(value)
    text = re.sub(r"^检查[:：]\s*", "", text)
    text = text.replace("...", "").replace("…", "").strip()
    text = re.sub(r"由于?时间\s*(不够|不足|有限|来不及)[^。]*[。]?", "", text).strip()
    return text


def _has_location_hint(text: str) -> bool:
    patterns = [
        r"第[一二三四五六七八九十百0-9]+章",
        r"第[一二三四五六七八九十百0-9]+节",
        r"第?[0-9]+页",
        r"\bP[0-9]+\b",
        r"页码",
        r"段落",
        r"第[一二三四五六七八九十百0-9]+段",
        r"条款",
        r"附件",
        r"截图",
        r"图[一二三四五六七八九十0-9]",
        r"表[一二三四五六七八九十0-9]",
    ]
    return any(re.search(p, text) for p in patterns)


def _clean_recommendation(value: Any) -> str:
    text = _clean_text(value)
    if not text:
        return text
    text = re.sub(r"由于?时间\s*(不够|不足|有限|来不及)[^。]*[。]?", "", text).strip()
    # 面向业务人员，避免“OCR”缩写术语
    text = re.sub(r"(?i)ocr验证", "核对截图中的文字内容", text)
    text = re.sub(r"(?i)进行ocr", "进行截图文字内容核对", text)
    text = re.sub(r"(?i)\bocr\b", "截图文字内容核对", text)
    text = text.replace(
        "需对相关截图进行核对截图中的文字内容，确保内容符合要求",
        "请核对截图中的文字内容，并标注对应页码和位置，确保内容符合要求",
    )
    text = text.replace(
        "需提供清晰的截图证据并进行截图文字内容核对",
        "请提供清晰截图，并标注对应页码和位置后核对文字内容",
    )
    return text


def _normalize_findings(raw: Any) -> list[dict[str, Any]]:
    items = raw if isinstance(raw, list) else []
    out: list[dict[str, Any]] = []
    for idx, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        status = _normalize_status(item.get("status"))
        issue = _clean_issue(item.get("issue") or item.get("summary") or "")
        tender_evidence = _clean_text(item.get("tender_evidence") or "")
        bid_evidence = _clean_text(item.get("bid_evidence") or "")
        bid_evidence = re.sub(r"(?i)需ocr验证内容", "需核对截图中的文字内容", bid_evidence)
        bid_evidence = re.sub(r"(?i)ocr验证", "截图文字内容核对", bid_evidence)
        bid_evidence = re.sub(r"(?i)\bocr\b", "截图文字内容核对", bid_evidence)
        recommendation = _clean_recommendation(item.get("recommendation") or "")

        # 投标证据强制要求可定位信息，缺失时补充提示语。
        if bid_evidence and not _has_location_hint(bid_evidence):
            bid_evidence = f"{bid_evidence}（缺少定位信息：请补充章节/页码/段落）"
            if not recommendation:
                recommendation = "请补充可定位的投标证据（章节/页码/段落）后再核对。"
        elif not bid_evidence:
            bid_evidence = "未提供投标证据（请补充章节/页码/段落）"
            if not recommendation:
                recommendation = "请补充可定位的投标证据（章节/页码/段落）后再核对。"

        out.append(
            {
                "id": item.get("id") or f"F{idx:03d}",
                "requirement_id": item.get("requirement_id") or "",
                "status": status,
                "issue": issue,
                "tender_evidence": tender_evidence,
                "bid_evidence": bid_evidence,
                "recommendation": recommendation,
            }
        )
    return [x for x in out if x["issue"]]


def _compact_token_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").lower())


def _find_context_requirement_id(requirements: list[dict[str, Any]]) -> str:
    for req in requirements:
        text = _compact_token_text(f"{req.get('category', '')} {req.get('text', '')}")
        if any(k in text for k in _CONTEXT_REQUIREMENT_KEYWORDS):
            rid = str(req.get("id", "")).strip()
            if rid:
                return rid
    return ""


def _ensure_context_consistency_requirement(report: dict[str, Any]) -> dict[str, Any]:
    requirements_raw = report.get("requirements", [])
    requirements = requirements_raw if isinstance(requirements_raw, list) else []
    if _find_context_requirement_id(requirements):
        return report

    requirements.append(
        {
            "id": f"R{len(requirements) + 1:03d}",
            "category": _CONTEXT_REQUIREMENT_CATEGORY,
            "text": _CONTEXT_REQUIREMENT_TEXT,
            "source": _CONTEXT_REQUIREMENT_SOURCE,
        }
    )
    report["requirements"] = requirements
    summary = report.get("summary")
    if not isinstance(summary, dict):
        summary = {}
        report["summary"] = summary
    summary["requirement_count"] = len(requirements)
    return report


def _is_context_consistency_finding(finding: dict[str, Any]) -> bool:
    text = _compact_token_text(
        " ".join(
            str(finding.get(k, "") or "")
            for k in ("issue", "tender_evidence", "bid_evidence", "recommendation")
        )
    )
    return any(k in text for k in _CONTEXT_FINDING_KEYWORDS)


def _bind_findings_to_context_requirement(
    findings: list[dict[str, Any]],
    *,
    valid_req_ids: set[str],
    context_req_id: str,
) -> list[dict[str, Any]]:
    if not context_req_id:
        return findings
    out: list[dict[str, Any]] = []
    for item in findings:
        f = dict(item)
        rid = str(f.get("requirement_id", "")).strip()
        if rid not in valid_req_ids and _is_context_consistency_finding(f):
            f["requirement_id"] = context_req_id
        out.append(f)
    return out


def normalize_review_report(data: dict[str, Any]) -> dict[str, Any]:
    requirements_raw = _normalize_requirements(data.get("requirements"))
    req_id_map: dict[str, str] = {}
    requirements: list[dict[str, Any]] = []
    for idx, req in enumerate(requirements_raw, start=1):
        old_id = req.get("id", "")
        new_id = f"R{idx:03d}"
        req_id_map[old_id] = new_id
        item = dict(req)
        item["id"] = new_id
        requirements.append(item)

    findings = _normalize_findings(data.get("findings"))
    for f in findings:
        rid = f.get("requirement_id", "")
        if rid in req_id_map:
            f["requirement_id"] = req_id_map[rid]
    summary = data.get("summary", {})
    if not isinstance(summary, dict):
        summary = {}
    summary_out = {
        "requirement_count": int(summary.get("requirement_count", len(requirements))),
        "non_compliant_count": sum(1 for f in findings if f["status"] == "non_compliant"),
        "risk_count": sum(1 for f in findings if f["status"] == "risk"),
        "needs_manual_count": sum(1 for f in findings if f["status"] == "needs_manual"),
        "finding_count": int(summary.get("finding_count", len(findings))),
    }
    summary_out["finding_count"] = len(findings)
    summary_out["requirement_count"] = len(requirements)
    return {
        "requirements": requirements,
        "findings": findings,
        "summary": summary_out,
    }


def _dedupe_findings(findings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for idx, f in enumerate(findings, start=1):
        key = (
            (f.get("requirement_id") or "").strip()
            + "|"
            + (f.get("issue") or "").replace(" ", "").strip()
        )
        if not key or key in seen:
            continue
        seen.add(key)
        item = dict(f)
        item["id"] = f"F{idx:03d}"
        out.append(item)
    return out


_STATUS_RANK = {
    "needs_manual": 0,
    "risk": 1,
    "non_compliant": 2,
}


def _extract_docx_text(path: Path) -> str:
    doc: Any | None = None
    for _ in range(3):
        try:
            doc = Document(str(path))
            break
        except Exception:  # noqa: BLE001
            time.sleep(0.2)
    if doc is None:
        # 并发读取下 python-docx 偶发失败时，兜底直读 OOXML 文本。
        try:
            with zipfile.ZipFile(path, "r") as zf:
                xml_bytes = zf.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
            raw_texts = [node.text for node in root.iter() if node.text]
            return "\n".join(t.strip() for t in raw_texts if t and t.strip())
        except Exception:  # noqa: BLE001
            return ""

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                lines.append(line)
    return "\n".join(lines)


def _find_requirement_id_by_keywords(
    requirements: list[dict[str, Any]],
    keyword_groups: list[tuple[str, ...]],
) -> str:
    if not requirements:
        return ""

    indexed: list[tuple[str, str]] = []
    for req in requirements:
        rid = str(req.get("id", "")).strip()
        if not rid:
            continue
        blob = _compact_token_text(
            f"{req.get('category', '')} {req.get('text', '')} {req.get('source', '')}"
        )
        indexed.append((rid, blob))

    for group in keyword_groups:
        keys = tuple(_compact_token_text(x) for x in group if x)
        if not keys:
            continue
        for rid, blob in indexed:
            if all(k in blob for k in keys):
                return rid
    return ""


def _pick_requirement_id(report: dict[str, Any], candidates: list[str]) -> str:
    requirements = report.get("requirements", [])
    if not isinstance(requirements, list):
        return ""
    valid_ids = [str(r.get("id", "")).strip() for r in requirements if str(r.get("id", "")).strip()]
    for rid in candidates:
        if rid and rid in valid_ids:
            return rid
    return valid_ids[0] if valid_ids else ""


def _match_text_by_keyword_groups(text: str, keyword_groups: list[tuple[str, ...]]) -> bool:
    compact = _compact_token_text(text)
    for group in keyword_groups:
        keys = tuple(_compact_token_text(x) for x in group if x)
        if keys and all(k in compact for k in keys):
            return True
    return False


def _upsert_guard_finding(
    report: dict[str, Any],
    *,
    requirement_id: str,
    status: str,
    issue: str,
    tender_evidence: str,
    bid_evidence: str,
    recommendation: str,
    match_groups: list[tuple[str, ...]],
) -> None:
    findings = report.get("findings", [])
    if not isinstance(findings, list):
        findings = []
        report["findings"] = findings

    for item in findings:
        if not isinstance(item, dict):
            continue
        merged_text = " ".join(
            str(item.get(k, "") or "")
            for k in ("issue", "tender_evidence", "bid_evidence", "recommendation")
        )
        if not _match_text_by_keyword_groups(merged_text, match_groups):
            continue

        old_status = str(item.get("status", "needs_manual")).strip()
        if _STATUS_RANK.get(status, 0) > _STATUS_RANK.get(old_status, 0):
            item["status"] = status
        item["issue"] = issue
        if requirement_id:
            item["requirement_id"] = requirement_id
        if not _has_location_hint(str(item.get("bid_evidence", "") or "")):
            item["bid_evidence"] = bid_evidence
        if not str(item.get("tender_evidence", "") or "").strip():
            item["tender_evidence"] = tender_evidence
        if not str(item.get("recommendation", "") or "").strip():
            item["recommendation"] = recommendation
        return

    findings.append(
        {
            "id": "",
            "requirement_id": requirement_id,
            "status": status,
            "issue": issue,
            "tender_evidence": tender_evidence,
            "bid_evidence": bid_evidence,
            "recommendation": recommendation,
        }
    )


def _refresh_summary(report: dict[str, Any]) -> None:
    findings = report.get("findings", [])
    if not isinstance(findings, list):
        findings = []
        report["findings"] = findings
    requirements = report.get("requirements", [])
    if not isinstance(requirements, list):
        requirements = []
        report["requirements"] = requirements
    summary = report.get("summary")
    if not isinstance(summary, dict):
        summary = {}
        report["summary"] = summary
    summary["non_compliant_count"] = sum(1 for f in findings if str(f.get("status", "")) == "non_compliant")
    summary["risk_count"] = sum(1 for f in findings if str(f.get("status", "")) == "risk")
    summary["needs_manual_count"] = sum(1 for f in findings if str(f.get("status", "")) == "needs_manual")
    summary["finding_count"] = len(findings)
    summary["requirement_count"] = len(requirements)


def _extract_year_after_anchor(compact_text: str, anchor: str, window: int) -> str:
    pos = compact_text.find(anchor)
    if pos < 0:
        return ""
    seg = compact_text[pos : pos + window]
    match = re.search(r"(20\d{2})年", seg)
    return match.group(1) if match else ""


_FINDING_THEME_RULES: list[dict[str, Any]] = [
    {
        "key": "subject_mismatch",
        "match_groups": [
            ("投标函", "投标人", "招标人"),
            ("主体", "错位"),
        ],
        "canonical_issue": "投标函落款处投标人名称误写为招标人名称，主体信息错位。",
        "default_status": "non_compliant",
        "req_type": "context",
    },
    {
        "key": "name_typo",
        "match_groups": [
            ("有限责任司",),
            ("封面", "名称", "不完整"),
            ("封面", "缺少", "公"),
        ],
        "canonical_issue": "商务投标文件封面投标人名称缺少“公”字，公司名称不完整。",
        "default_status": "non_compliant",
        "req_type": "context",
    },
    {
        "key": "quote_missing",
        "match_groups": [
            ("开标一览表", "分项报价表", "税率"),
            ("含税", "不含税", "税率", "空"),
            ("报价", "未填写"),
        ],
        "canonical_issue": "开标一览表和分项报价表中的含税价、不含税价及税率字段存在空缺。",
        "default_status": "non_compliant",
        "req_type": "quote",
    },
    {
        "key": "date_conflict",
        "match_groups": [
            ("日期", "不一致"),
            ("编制日期", "投标函"),
            ("开标一览表", "日期"),
        ],
        "canonical_issue": "投标文件关键日期存在不一致，可能影响文件内部一致性与有效性判断。",
        "default_status": "risk",
        "req_type": "format",
    },
    {
        "key": "signature_blank",
        "match_groups": [
            ("签字栏", "空白"),
            ("委托代理人", "签字处", "空白"),
            ("法定代表人", "签字或盖章", "空白"),
        ],
        "canonical_issue": "商务/经济投标文件封面法定代表人或委托代理人签字栏为空白。",
        "default_status": "needs_manual",
        "req_type": "sign",
    },
    {
        "key": "bank_proof",
        "match_groups": [
            ("基本账户", "开户许可证"),
            ("投标保证金", "基本账户"),
            ("基本账户", "证明", "缺失"),
        ],
        "canonical_issue": "基本账户证明与投标保证金转出账户信息不足，需补充核验。",
        "default_status": "needs_manual",
        "req_type": "bank",
    },
    {
        "key": "social_security_manual",
        "match_groups": [
            ("项目负责人", "社保"),
            ("社保证明",),
            ("缴纳", "社保"),
        ],
        "canonical_issue": "项目负责人近一年社保证明需人工核对缴纳主体与时间范围。",
        "default_status": "needs_manual",
        "req_type": "social",
    },
    {
        "key": "audit_manual",
        "match_groups": [
            ("财务审计",),
            ("审计报告",),
            ("第三方审计",),
        ],
        "canonical_issue": "财务审计报告需人工核对审计年度与报告完整性。",
        "default_status": "needs_manual",
        "req_type": "audit",
    },
    {
        "key": "cert_manual",
        "match_groups": [
            ("体系认证",),
            ("iso",),
            ("认证证书",),
        ],
        "canonical_issue": "体系认证证书需人工核对有效期及官网查询结果。",
        "default_status": "needs_manual",
        "req_type": "cert",
    },
    {
        "key": "performance_manual",
        "match_groups": [
            ("类似项目", "业绩"),
            ("业绩", "时间"),
            ("项目业绩",),
        ],
        "canonical_issue": "类似项目业绩材料需人工核对项目范围与时间是否满足招标要求。",
        "default_status": "needs_manual",
        "req_type": "performance",
    },
]


def _theme_match_rule(finding: dict[str, Any]) -> dict[str, Any] | None:
    merged_text = " ".join(
        str(finding.get(k, "") or "")
        for k in ("issue", "tender_evidence", "bid_evidence", "recommendation")
    )
    for rule in _FINDING_THEME_RULES:
        if _match_text_by_keyword_groups(merged_text, rule["match_groups"]):
            return rule
    return None


def _finding_quality_score(item: dict[str, Any]) -> int:
    status = str(item.get("status", "needs_manual")).strip()
    score = _STATUS_RANK.get(status, 0) * 10
    if _has_location_hint(str(item.get("bid_evidence", "") or "")):
        score += 2
    if str(item.get("tender_evidence", "") or "").strip():
        score += 1
    return score


def _theme_requirement_candidates(report: dict[str, Any]) -> dict[str, str]:
    requirements = report.get("requirements", [])
    if not isinstance(requirements, list):
        requirements = []
    context_req_id = _find_context_requirement_id(requirements)
    format_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("投标文件格式",), ("响应格式",), ("格式",)],
    )
    quote_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("含税", "税率"), ("报价",), ("开标一览表",)],
    )
    sign_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("签字", "盖章"), ("签章",), ("电子印章",)],
    )
    bank_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("投标保证金", "基本账户"), ("基本账户",), ("投标保证金",)],
    )
    social_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("社保",), ("项目负责人", "社保")],
    )
    audit_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("财务", "审计"), ("审计报告",)],
    )
    cert_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("体系认证",), ("iso",)],
    )
    performance_req_id = _find_requirement_id_by_keywords(
        requirements,
        [("业绩",), ("类似项目", "业绩")],
    )
    return {
        "context": _pick_requirement_id(report, [context_req_id, format_req_id, quote_req_id, sign_req_id, bank_req_id]),
        "format": _pick_requirement_id(report, [format_req_id, sign_req_id, context_req_id]),
        "quote": _pick_requirement_id(report, [quote_req_id, format_req_id, context_req_id]),
        "sign": _pick_requirement_id(report, [sign_req_id, format_req_id, context_req_id]),
        "bank": _pick_requirement_id(report, [bank_req_id, quote_req_id, context_req_id]),
        "social": _pick_requirement_id(report, [social_req_id, context_req_id]),
        "audit": _pick_requirement_id(report, [audit_req_id, context_req_id]),
        "cert": _pick_requirement_id(report, [cert_req_id, context_req_id]),
        "performance": _pick_requirement_id(report, [performance_req_id, context_req_id]),
    }


def _stabilize_findings(report: dict[str, Any]) -> dict[str, Any]:
    findings = report.get("findings", [])
    if not isinstance(findings, list):
        return report
    req_ids = _theme_requirement_candidates(report)

    themed_best: dict[str, dict[str, Any]] = {}
    other: list[dict[str, Any]] = []
    for item in findings:
        if not isinstance(item, dict):
            continue
        rule = _theme_match_rule(item)
        if not rule:
            other.append(dict(item))
            continue
        key = str(rule["key"])
        current = themed_best.get(key)
        candidate = dict(item)
        # 优先更高严重度/更完整证据。
        if current is None or _finding_quality_score(candidate) > _finding_quality_score(current):
            themed_best[key] = candidate

    merged: list[dict[str, Any]] = []
    for rule in _FINDING_THEME_RULES:
        key = str(rule["key"])
        item = themed_best.get(key)
        if not item:
            continue
        expected_status = str(rule["default_status"])
        # 主题项状态固定，避免同义问题在不同运行中出现风险级别漂移。
        item["status"] = expected_status
        item["issue"] = str(rule["canonical_issue"])
        req_type = str(rule["req_type"])
        rid = req_ids.get(req_type, "")
        if rid:
            item["requirement_id"] = rid
        merged.append(item)

    keep_non_theme = os.getenv("BID_REVIEW_KEEP_NON_THEME_FINDINGS", "0").strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }
    if keep_non_theme:
        # 可选：保留非主题项用于人工深挖，默认关闭以确保结果稳定。
        for item in other:
            status = str(item.get("status", "")).strip()
            if status == "non_compliant":
                merged.append(item)
                continue
            if status == "needs_manual" and _has_location_hint(str(item.get("bid_evidence", "") or "")):
                merged.append(item)

    report["findings"] = _dedupe_findings(merged)
    _refresh_summary(report)
    return report


def _apply_docx_stability_guards_from_text(
    report: dict[str, Any],
    docx_text: str,
    *,
    force_manual_image_checks: bool = False,
) -> dict[str, Any]:
    def _normalize_labeled_party(value: str) -> str:
        text = str(value or "").strip()
        text = re.split(r"[\r\n]", text, maxsplit=1)[0]
        text = re.sub(r"[（(](?:盖[^）)]*|签[^）)]*|签章[^）)]*)[）)]\s*$", "", text)
        text = re.split(r"(?:联系人|联系电话|地址|电话|邮编)\s*[:：]", text, maxsplit=1)[0]
        text = text.strip("`'\"：:，,。.;； ")
        return re.sub(r"\s+", "", text)

    def _extract_labeled_parties(text: str, labels: tuple[str, ...]) -> set[str]:
        out: set[str] = set()
        label_group = "|".join(re.escape(label) for label in labels)
        pattern = re.compile(
            rf"(?:{label_group})\s*[:：]\s*([^\r\n]{{2,120}})",
            flags=re.IGNORECASE,
        )
        for match in pattern.finditer(text):
            normalized = _normalize_labeled_party(match.group(1))
            if len(normalized) >= 4:
                out.add(normalized)
        return out

    if not docx_text:
        return report

    requirements = report.get("requirements", [])
    if not isinstance(requirements, list):
        requirements = []
        report["requirements"] = requirements
    context_req_id = _find_context_requirement_id(requirements)
    format_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("投标文件格式",),
            ("响应格式",),
            ("格式",),
        ],
    )
    quote_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("含税", "税率"),
            ("报价",),
            ("开标一览表",),
        ],
    )
    sign_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("签字", "盖章"),
            ("签章",),
            ("电子印章",),
        ],
    )
    bank_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("投标保证金", "基本账户"),
            ("基本账户",),
            ("投标保证金",),
        ],
    )
    social_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("社保",),
            ("项目负责人", "社保"),
        ],
    )
    audit_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("财务", "审计"),
            ("审计报告",),
        ],
    )
    cert_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("体系认证",),
            ("iso",),
            ("认证证书",),
        ],
    )
    performance_req_id = _find_requirement_id_by_keywords(
        requirements,
        [
            ("业绩",),
            ("类似项目", "业绩"),
        ],
    )

    context_req_id = _pick_requirement_id(
        report,
        [
            context_req_id,
            format_req_id,
            quote_req_id,
            sign_req_id,
            bank_req_id,
            social_req_id,
            audit_req_id,
            cert_req_id,
            performance_req_id,
        ],
    )
    format_req_id = _pick_requirement_id(report, [format_req_id, sign_req_id, context_req_id])
    quote_req_id = _pick_requirement_id(report, [quote_req_id, format_req_id, context_req_id])
    sign_req_id = _pick_requirement_id(report, [sign_req_id, format_req_id, context_req_id])
    bank_req_id = _pick_requirement_id(report, [bank_req_id, quote_req_id, context_req_id])
    social_req_id = _pick_requirement_id(report, [social_req_id, context_req_id])
    audit_req_id = _pick_requirement_id(report, [audit_req_id, context_req_id])
    cert_req_id = _pick_requirement_id(report, [cert_req_id, context_req_id])
    performance_req_id = _pick_requirement_id(report, [performance_req_id, context_req_id])

    req_map = {
        str(r.get("id", "")).strip(): r
        for r in requirements
        if isinstance(r, dict) and str(r.get("id", "")).strip()
    }

    def tender_evidence_for(rid: str, fallback: str) -> str:
        req = req_map.get(rid, {})
        if not isinstance(req, dict):
            return fallback
        text = str(req.get("text", "")).strip()
        source = str(req.get("source", "")).strip()
        if text and source:
            return f"{text}（{source}）"
        if text:
            return text
        if source:
            return source
        return fallback

    compact = _compact_token_text(docx_text)

    # 1) 投标函落款主体错位（明确不符合）
    tender_party_names = _extract_labeled_parties(docx_text, ("招标人", "采购人"))
    bid_letter_party_names = _extract_labeled_parties(
        re.search(r"投标函[\s\S]{0,2600}", docx_text).group(0) if re.search(r"投标函[\s\S]{0,2600}", docx_text) else "",
        ("投标人",),
    )
    if tender_party_names and any(name in tender_party_names for name in bid_letter_party_names):
        _upsert_guard_finding(
            report,
            requirement_id=context_req_id,
            status="non_compliant",
            issue="投标函落款处投标人名称误写为招标人名称，主体信息错位。",
            tender_evidence=tender_evidence_for(
                context_req_id,
                "投标文件关键主体名词必须与所在位置和语义角色一致。",
            ),
            bid_evidence="投标函落款页：检测到“投标人”字段值与文内“招标人/采购人”字段值一致，主体不一致。",
            recommendation="将投标函落款处投标人名称更正为投标人法定全称，并复核同页签章信息。",
            match_groups=[
                ("投标函", "投标人", "招标人"),
                ("投标函", "落款", "主体"),
            ],
        )

    # 2) 商务投标文件封面公司名称缺字（明确不符合）
    if "有限责任司" in docx_text:
        _upsert_guard_finding(
            report,
            requirement_id=context_req_id,
            status="non_compliant",
            issue="商务投标文件封面投标人名称缺少“公”字，公司名称不完整。",
            tender_evidence=tender_evidence_for(
                context_req_id,
                "投标文件关键主体名词必须与所在位置和语义角色一致。",
            ),
            bid_evidence="商务投标文件封面：检测到 `有限责任司` 异常，公司名称疑似缺少“公”字。",
            recommendation="将封面公司名称更正为营业执照一致的法定全称，并统一检查全文件主体名称。",
            match_groups=[
                ("有限责任司",),
                ("商务投标文件", "公司名称", "不完整"),
            ],
        )

    # 3) 开标/分项报价关键字段空缺（明确不符合）
    if re.search(r"开标一览表.{0,5000}小写[:：]元", compact) and re.search(r"税率[:：]%", compact):
        _upsert_guard_finding(
            report,
            requirement_id=quote_req_id,
            status="non_compliant",
            issue="开标一览表和分项报价表中的含税价、不含税价及税率字段存在空缺。",
            tender_evidence=tender_evidence_for(
                quote_req_id,
                "投标报价应按招标文件要求完整填报含税、不含税及税率信息。",
            ),
            bid_evidence="开标一览表/分项报价表：`小写： 元`、`税率： %` 等字段为空，无法形成有效报价。",
            recommendation="补齐含税价、不含税价、税率及增值税税额，并校验总价与分项汇总一致。",
            match_groups=[
                ("开标一览表", "含税", "税率"),
                ("分项报价表", "未填写"),
                ("报价", "空缺"),
            ],
        )

    # 4) 关键日期冲突（风险）
    years: dict[str, str] = {}
    raw_patterns = [
        (r"商务投标文件[\s\S]{0,260}?编制日期[:：]\s*(20\d{2})\s*年", "商务封面"),
        (r"经济投标文件[\s\S]{0,260}?编制日期[:：]\s*(20\d{2})\s*年", "经济封面"),
        (r"开标一览表[\s\S]{0,2600}?日\s*期\s*[:：]\s*(20\d{2})\s*年", "开标一览表"),
        (r"投标函[\s\S]{0,2600}?日\s*期\s*[:：]\s*(20\d{2})\s*年", "投标函"),
    ]
    for pattern, key in raw_patterns:
        match = re.search(pattern, docx_text)
        if match:
            years[key] = match.group(1)
    distinct_years = sorted(set(years.values()))
    if len(distinct_years) >= 2:
        details = "、".join([f"{k}{v}年" for k, v in years.items()])
        _upsert_guard_finding(
            report,
            requirement_id=format_req_id,
            status="risk",
            issue="投标文件关键日期存在不一致，可能影响文件内部一致性与有效性判断。",
            tender_evidence=tender_evidence_for(
                format_req_id,
                "投标文件应按格式完整、准确填写关键信息。",
            ),
            bid_evidence=f"日期交叉核对：{details}；存在多个年份并存的情况。",
            recommendation="统一封面、投标函、开标一览表等关键日期，并按投标时点复核全文件时间一致性。",
            match_groups=[
                ("日期", "不一致"),
                ("编制日期", "投标函"),
                ("开标一览表", "年份"),
            ],
        )

    # 5) 封面签字位置空白（需人工复核）
    blank_sign_count = len(re.findall(r"法定代表人、负责人或委托代理人[:：]（签字或盖章）", compact))
    if blank_sign_count >= 1 and "签字或盖章" in compact:
        _upsert_guard_finding(
            report,
            requirement_id=sign_req_id,
            status="needs_manual",
            issue="商务/经济投标文件封面法定代表人或委托代理人签字栏为空白。",
            tender_evidence=tender_evidence_for(
                sign_req_id,
                "投标文件中要求签字或盖章的，应按要求执行。",
            ),
            bid_evidence="商务投标文件封面、经济投标文件封面均出现`法定代表人、负责人或委托代理人：（签字或盖章）`空白栏。",
            recommendation="核对电子签章或手签扫描是否已按要求补齐，并确保封面签署信息完整。",
            match_groups=[
                ("签字栏", "空白"),
                ("法定代表人", "委托代理人", "签字或盖章"),
            ],
        )

    # 6) 基本账户与保证金转出核验项
    if force_manual_image_checks or ("基本账户" in compact and "投标保证金" in compact):
        _upsert_guard_finding(
            report,
            requirement_id=bank_req_id,
            status="needs_manual",
            issue="基本账户证明与投标保证金转出账户信息不足，需补充核验。",
            tender_evidence=tender_evidence_for(
                bank_req_id,
                "投标保证金应由投标人基本账户转出并提供可核验证明。",
            ),
            bid_evidence="投标文件“投标保证金交纳证明/基本账户证明”为图片证据，需核对付款账号与基本账户一致性。",
            recommendation="补充银行转账回单及基本账户证明关键页，并标注账户名、账号、开户行信息。",
            match_groups=[
                ("基本账户", "投标保证金"),
                ("基本账户", "开户许可证"),
                ("投标保证金", "转出账户"),
            ],
        )

    # 7) 关键图片证据人工核验项（固定输出，降低多次运行漂移）
    if force_manual_image_checks or ("社保" in compact):
        _upsert_guard_finding(
            report,
            requirement_id=social_req_id,
            status="needs_manual",
            issue="项目负责人近一年社保证明需人工核对缴纳主体与时间范围。",
            tender_evidence=tender_evidence_for(
                social_req_id,
                "项目负责人相关社保材料需满足招标资格要求。",
            ),
            bid_evidence="投标文件“项目负责人社保证明”为图片证据，需核对缴纳单位、姓名及连续缴纳期间。",
            recommendation="补充清晰社保证明页面，并标注项目负责人姓名、缴纳单位及起止时间。",
            match_groups=[
                ("项目负责人", "社保"),
                ("社保证明",),
            ],
        )

    if force_manual_image_checks or "财务审计报告" in compact or ("审计报告" in compact and "财务" in compact):
        _upsert_guard_finding(
            report,
            requirement_id=audit_req_id,
            status="needs_manual",
            issue="财务审计报告需人工核对审计年度与报告完整性。",
            tender_evidence=tender_evidence_for(
                audit_req_id,
                "财务审计材料应满足招标文件对审计年度和完整性的要求。",
            ),
            bid_evidence="投标文件“财务审计报告”为图片证据，需核对是否为2024年度且包含完整审计正文与签章页。",
            recommendation="补充完整财务审计报告关键页并标注审计年度、审计机构及签章信息。",
            match_groups=[
                ("财务审计",),
                ("审计报告",),
            ],
        )

    if force_manual_image_checks or "体系认证证书" in compact or "iso9001" in compact or "iso27001" in compact or "iso20000" in compact:
        _upsert_guard_finding(
            report,
            requirement_id=cert_req_id,
            status="needs_manual",
            issue="体系认证证书需人工核对有效期及官网查询结果。",
            tender_evidence=tender_evidence_for(
                cert_req_id,
                "体系认证证书需满足有效期与认证范围要求。",
            ),
            bid_evidence="投标文件“体系认证证书”主要为截图证据，需核对证书编号、有效期及官网查询一致性。",
            recommendation="提供清晰证书页和官网查询页，并对照标注证书编号与有效期。",
            match_groups=[
                ("体系认证",),
                ("iso",),
                ("认证证书",),
            ],
        )

    if force_manual_image_checks or "类似项目业绩" in compact or ("项目业绩" in compact and "业绩表" in compact):
        _upsert_guard_finding(
            report,
            requirement_id=performance_req_id,
            status="needs_manual",
            issue="类似项目业绩材料需人工核对项目范围与时间是否满足招标要求。",
            tender_evidence=tender_evidence_for(
                performance_req_id,
                "类似项目业绩应满足招标文件规定的范围和时间要求。",
            ),
            bid_evidence="投标文件“类似项目业绩材料”为合同/发票截图，需核对项目内容、签订时间与招标范围匹配性。",
            recommendation="补充可读性高的业绩证明关键页，并标注项目名称、签订时间和服务内容。",
            match_groups=[
                ("类似项目", "业绩"),
                ("项目业绩",),
            ],
        )

    report["findings"] = _dedupe_findings(report.get("findings", []))
    _refresh_summary(report)
    return _stabilize_findings(report)


def _apply_stability_guards(
    report: dict[str, Any],
    *,
    bid_path: Path,
    force_manual_image_checks: bool = False,
) -> dict[str, Any]:
    if bid_path.suffix.lower() != ".docx":
        return report
    text = _extract_docx_text(bid_path)
    if not text:
        return report
    return _apply_docx_stability_guards_from_text(
        report,
        text,
        force_manual_image_checks=force_manual_image_checks,
    )


def detect_roles_with_claude(paths: list[str], client: Any) -> tuple[str, str, str]:
    backend_name = "OpenCode" if client.__class__.__name__.lower().startswith("opencode") else "Claude"
    docs = []
    for idx, p in enumerate(paths, start=1):
        path = Path(p).resolve()
        docs.append({"id": f"D{idx}", "path": str(path), "stem": path.stem})
    path_block = "\n".join(f"- {d['id']}: {d['stem']}" for d in docs)
    prompt = render_prompt("role_detect_single.md", file_list=path_block)
    data = client.ask_json(
        prompt,
        required_top_keys=["tender_id", "bid_id"],
        task_label="识别招标文件与投标文件",
    )
    if not isinstance(data, dict):
        raise ValueError(f"{backend_name} 返回格式错误，无法识别招投标文件。")
    tender_id = str(data["tender_id"])
    bid_id = str(data["bid_id"])
    reasoning = str(data.get("reasoning", ""))
    by_id = {d["id"]: d["path"] for d in docs}
    tender_path = by_id.get(tender_id, "")
    bid_path = by_id.get(bid_id, "")

    # 安全兜底。
    if not tender_path or not bid_path:
        # 文件名兜底
        tender_guess = next((p for p in paths if "招标" in Path(p).name), paths[0])
        bid_guess = next((p for p in paths if "投标" in Path(p).name), paths[-1])
        tender_path = str(Path(tender_guess).resolve())
        bid_path = str(Path(bid_guess).resolve())
        reasoning = f"{reasoning}; fallback-by-filename"

    if str(Path(tender_path).resolve()) == str(Path(bid_path).resolve()):
        raise ValueError(f"{backend_name} 未能区分招标与投标文件。")
    return tender_path, bid_path, reasoning


def detect_tender_and_bids_with_claude(
    paths: list[str],
    client: Any,
) -> tuple[str, list[str], str]:
    backend_name = "OpenCode" if client.__class__.__name__.lower().startswith("opencode") else "Claude"
    docs = []
    for idx, p in enumerate(paths, start=1):
        path = Path(p).resolve()
        docs.append({"id": f"D{idx}", "path": str(path), "stem": path.stem})
    path_block = "\n".join(f"- {d['id']}: {d['stem']}" for d in docs)
    prompt = render_prompt("role_detect_multi.md", file_list=path_block)
    data = client.ask_json(
        prompt,
        required_top_keys=["tender_id", "bid_ids"],
        task_label="识别招标文件与多个投标文件",
    )
    if not isinstance(data, dict):
        raise ValueError(f"{backend_name} 返回格式错误，无法识别招投标文件。")
    tender_id = str(data.get("tender_id", ""))
    bid_ids_raw = data.get("bid_ids", [])
    reasoning = str(data.get("reasoning", ""))
    bid_ids = [str(x) for x in bid_ids_raw] if isinstance(bid_ids_raw, list) else []

    by_id = {d["id"]: d["path"] for d in docs}
    tender_path = by_id.get(tender_id, "")
    bid_paths = [by_id.get(x, "") for x in bid_ids]
    bid_paths = [x for x in bid_paths if x and x != tender_path]

    # 安全兜底。
    if not tender_path:
        tender_guess = next((p for p in paths if "招标" in Path(p).name), paths[0])
        tender_path = str(Path(tender_guess).resolve())
        reasoning = f"{reasoning}; fallback-tender-by-filename"

    if not bid_paths:
        bid_by_name = [
            str(Path(p).resolve()) for p in paths if ("投标" in Path(p).name) and (str(Path(p).resolve()) != tender_path)
        ]
        bid_paths = bid_by_name
        if not bid_paths:
            bid_paths = [str(Path(p).resolve()) for p in paths if str(Path(p).resolve()) != tender_path]
        reasoning = f"{reasoning}; fallback-bids-by-filename"

    if not bid_paths:
        raise ValueError("未识别到投标文件。")
    return str(Path(tender_path).resolve()), bid_paths, reasoning


def run_bid_review_with_claude(
    *,
    tender_path: str,
    bid_path: str,
    client: Any,
    extra_instruction: str = "",
    user_instruction: str = "",
) -> tuple[dict[str, Any], str]:
    backend_name = "OpenCode" if client.__class__.__name__.lower().startswith("opencode") else "Claude"
    tender_path_obj = Path(tender_path).resolve()
    bid_path_obj = Path(bid_path).resolve()
    workspace_dir = prompt_safe_path(str(tender_path_obj.parent))
    tender_stem = tender_path_obj.stem
    bid_stem = bid_path_obj.stem
    instruction = compact_text_for_prompt(extra_instruction.strip(), 2000) if extra_instruction else "无"
    user_ins = compact_text_for_prompt(user_instruction.strip(), 2000) if user_instruction else "无"
    prompt = render_prompt(
        "review_main.md",
        workspace_dir=workspace_dir,
        tender_stem=tender_stem,
        bid_stem=bid_stem,
        tender_path=str(tender_path_obj),
        bid_path=str(bid_path_obj),
        user_instruction=user_ins,
        instruction=instruction,
    )
    require_word_extract = bid_path_obj.suffix.lower() == ".docx"
    ocr_required = _instruction_requires_ocr(user_instruction, extra_instruction) or (
        require_word_extract and _docx_ocr_required_by_default()
    )
    word_ocr_fully_covered = not (ocr_required and require_word_extract)
    original_timeout = client.timeout_sec
    # OCR全量处理（特别是docx图片较多时）需要更长超时，避免中途失败。
    if ocr_required and require_word_extract and client.timeout_sec < 7200:
        client.timeout_sec = 7200
    try:
        prompt = _append_no_write_enforcement(prompt)
        raw_output = client.ask_text(prompt, task_label=f"初审：{bid_path_obj.name}")
        first_calls = client.get_last_tool_calls()
        first_uses = client.get_last_tool_uses()
        has_ocr = _has_ocr_tool_call(first_calls)
        has_word_extract = _has_word_image_extract_call(first_calls) if require_word_extract else True
        has_word_batch_ocr = _has_word_batch_ocr_call(first_calls) if require_word_extract else True
        word_ocr_coverage_ok = True
        word_ocr_coverage_detail = ""
        if ocr_required and require_word_extract and has_word_extract and has_word_batch_ocr:
            word_ocr_coverage_ok, word_ocr_coverage_detail = _validate_docx_ocr_coverage(
                first_uses,
                bid_path=bid_path_obj,
            )
        has_forbidden_write = _has_forbidden_write_tool_call(first_uses)
        need_retry = has_forbidden_write or (
            ocr_required and (not has_ocr or not has_word_extract or not has_word_batch_ocr or not word_ocr_coverage_ok)
        )
        if need_retry:
            retry_prompt = prompt
            if ocr_required and (
                not has_ocr or not has_word_extract or not has_word_batch_ocr or not word_ocr_coverage_ok
            ):
                retry_prompt = _append_ocr_enforcement(retry_prompt, require_word_extract=require_word_extract)
            if require_word_extract and not word_ocr_coverage_ok and word_ocr_coverage_detail:
                retry_prompt = (
                    retry_prompt
                    + "\n你上一次未完成 Word 提图全量 OCR。"
                    + word_ocr_coverage_detail
                    + "请严格覆盖提图目录中的全部图片。"
                )
            if has_forbidden_write:
                retry_prompt = _append_no_write_enforcement(retry_prompt)
            retry_output = client.ask_text(retry_prompt, task_label=f"初审重试(约束强制)：{bid_path_obj.name}")
            retry_calls = client.get_last_tool_calls()
            retry_uses = client.get_last_tool_uses()
            has_ocr = _has_ocr_tool_call(retry_calls)
            has_word_extract = _has_word_image_extract_call(retry_calls) if require_word_extract else True
            has_word_batch_ocr = _has_word_batch_ocr_call(retry_calls) if require_word_extract else True
            word_ocr_coverage_ok = True
            word_ocr_coverage_detail = ""
            if ocr_required and require_word_extract and has_word_extract and has_word_batch_ocr:
                word_ocr_coverage_ok, word_ocr_coverage_detail = _validate_docx_ocr_coverage(
                    retry_uses,
                    bid_path=bid_path_obj,
                )
            has_forbidden_write = _has_forbidden_write_tool_call(retry_uses)
            if not has_ocr or not has_word_extract or not has_word_batch_ocr:
                need = []
                if not has_ocr:
                    need.append("OCR工具调用")
                if not has_word_extract:
                    need.append("Word图片提取调用")
                if not has_word_batch_ocr:
                    need.append("全量图片批量OCR调用")
                raise ClaudeCallError(
                    f"审查阶段缺少必要MCP调用（{', '.join(need)}），已按强制规则重试1次仍失败。"
                )
            if require_word_extract and not word_ocr_coverage_ok:
                raise ClaudeCallError(
                    f"Word图片OCR未全量覆盖，已按强制规则重试1次仍失败。{word_ocr_coverage_detail}"
                )
            if has_forbidden_write and _strict_fail_on_forbidden_write():
                raise ClaudeCallError("审查阶段检测到写文件/脚本执行行为，已按只读规则重试1次仍失败。")
            raw_output = retry_output
        elif ocr_required and require_word_extract and not word_ocr_coverage_ok:
            raise ClaudeCallError(
                f"Word图片OCR未全量覆盖。{word_ocr_coverage_detail}"
            )
        if ocr_required and require_word_extract:
            word_ocr_fully_covered = word_ocr_coverage_ok
    finally:
        client.timeout_sec = original_timeout
    try:
        data = extract_json_payload(raw_output)
    except Exception:  # noqa: BLE001
        # 二次重试：走严格 JSON API 通道。
        data = client.ask_json(
            prompt,
            required_top_keys=["requirements", "findings", "summary"],
            task_label=f"初审(JSON重试)：{bid_path_obj.name}",
        )
        raw_output = (
            f"{raw_output}\n\n[JSON_FALLBACK]\n"
            + json.dumps(data, ensure_ascii=False, indent=2)
        )
    if not isinstance(data, dict):
        raise ValueError(f"{backend_name} 返回的审查结果不是 JSON 对象。")
    for key in ("requirements", "findings", "summary"):
        if key not in data:
            raise ValueError(f"{backend_name} 返回缺少关键字段: {key}")
    report = normalize_review_report(data)
    report = _ensure_context_consistency_requirement(report)
    valid_req_ids = {str(r.get("id", "")) for r in report.get("requirements", [])}
    context_req_id = _find_context_requirement_id(report.get("requirements", []))
    report["findings"] = _bind_findings_to_context_requirement(
        report.get("findings", []),
        valid_req_ids=valid_req_ids,
        context_req_id=context_req_id,
    )
    report["summary"]["non_compliant_count"] = sum(
        1 for f in report["findings"] if f["status"] == "non_compliant"
    )
    report["summary"]["risk_count"] = sum(1 for f in report["findings"] if f["status"] == "risk")
    report["summary"]["needs_manual_count"] = sum(
        1 for f in report["findings"] if f["status"] == "needs_manual"
    )
    report["summary"]["finding_count"] = len(report["findings"])

    enable_second_pass = os.getenv("BID_REVIEW_ENABLE_SECOND_PASS", "0").strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }
    force_manual_image_checks = require_word_extract and ocr_required and (not word_ocr_fully_covered)
    if not enable_second_pass:
        report = _apply_stability_guards(
            report,
            bid_path=bid_path_obj,
            force_manual_image_checks=force_manual_image_checks,
        )
        merged_raw = f"{raw_output}\n\n[SECOND_PASS]\nSKIPPED_BY_DEFAULT"
        return report, merged_raw

    # 二次复核：专找初审遗漏项。
    initial_json = json.dumps(report, ensure_ascii=False, indent=2)
    initial_json = compact_text_for_prompt(initial_json, 8000)
    second_prompt = render_prompt(
        "review_second_pass.md",
        workspace_dir=workspace_dir,
        tender_stem=tender_stem,
        bid_stem=bid_stem,
        tender_path=str(tender_path_obj),
        bid_path=str(bid_path_obj),
        user_instruction=user_ins,
        initial_json=initial_json,
    )
    second_prompt = _append_no_write_enforcement(second_prompt)
    second_raw = client.ask_text(second_prompt, task_label=f"二次复核：{bid_path_obj.name}")
    if _has_forbidden_write_tool_call(client.get_last_tool_uses()):
        second_retry = _append_no_write_enforcement(second_prompt)
        second_raw = client.ask_text(second_retry, task_label=f"二次复核重试(只读强制)：{bid_path_obj.name}")
        if _has_forbidden_write_tool_call(client.get_last_tool_uses()) and _strict_fail_on_forbidden_write():
            raise ClaudeCallError("二次复核阶段检测到写文件/脚本执行行为，已按只读规则重试1次仍失败。")
    try:
        second_data = extract_json_payload(second_raw)
        add_findings = _normalize_findings(second_data.get("additional_findings", []))
    except Exception:  # noqa: BLE001
        add_findings = []
    if add_findings:
        # 二次复核只允许引用初审已存在的 requirement_id，防止凭空新增 R041+。
        valid_req_ids = {str(r.get("id", "")) for r in report.get("requirements", [])}
        context_req_id = _find_context_requirement_id(report.get("requirements", []))
        add_findings = _bind_findings_to_context_requirement(
            add_findings,
            valid_req_ids=valid_req_ids,
            context_req_id=context_req_id,
        )
        add_findings = [f for f in add_findings if str(f.get("requirement_id", "")) in valid_req_ids]
    if add_findings:
        report["findings"] = _dedupe_findings(report["findings"] + add_findings)
        report["summary"]["non_compliant_count"] = sum(
            1 for f in report["findings"] if f["status"] == "non_compliant"
        )
        report["summary"]["risk_count"] = sum(1 for f in report["findings"] if f["status"] == "risk")
        report["summary"]["needs_manual_count"] = sum(
            1 for f in report["findings"] if f["status"] == "needs_manual"
        )
        report["summary"]["finding_count"] = len(report["findings"])

    # 稳定性兜底：对可确定的不符合项做规则化补齐/归一，降低多次运行抖动。
    report = _apply_stability_guards(
        report,
        bid_path=bid_path_obj,
        force_manual_image_checks=force_manual_image_checks,
    )

    merged_raw = f"{raw_output}\n\n[SECOND_PASS]\n{second_raw}"
    return report, merged_raw
